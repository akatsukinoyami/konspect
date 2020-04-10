import docx
import random
import os

class Conspect:

	def docx(self, msg, path):
		doc1 = docx.Document(path)
		doc2 = docx.Document()

		for paragraph in doc1.paragraphs:
			doc2para = doc2.add_paragraph()
			for run in paragraph.runs:
				for symbol in run.text:
					self.add_text(doc2para, symbol)
		doc2.save(path) 
		return path
	
	def txt(self, msg, path):
		f = open(path)
		doc2 = docx.Document()

		for line in f.readlines():
			doc2para = doc2.add_paragraph()
			for symbol in line:
				self.add_text(doc2para, symbol)

		path = os.path.abspath(msg.document.file_name.replace('.txt', '.docx'))
		doc2.save(path) 
		return path
	
	def text(self, msg):
		doc2 = docx.Document()

		doc2para = doc2.add_paragraph()
		for symbol in msg.text:
			self.add_text(doc2para, symbol)

		path = os.path.abspath(f'{msg.chat.id}_{msg.message_id}.docx')
		doc2.save(path) 
		return path
	
	def add_text(self, para, text):
		a 			= random.randint(1,4)
		run 		= para.add_run(text)
		font 		= run.font	
		font.name	= 'Ебать'	

		if   a == 2:	font.bold	= True
		elif a == 3:	font.italic	= True
		elif a == 4:
			font.bold	= True
			font.italic	= True