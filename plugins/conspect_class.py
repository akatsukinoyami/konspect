import docx, random, os

class Conspect:
	def docx(self, msg, path):
		doc1 = docx.Document(path)
		doc2 = docx.Document()

		for paragraph in doc1.paragraphs:
			doc2para = doc2.add_paragraph()
			for run in paragraph.runs:
				for symbol in run.text:
					self.add_text(doc2para, symbol)
		self.save_file(doc2, path)
		return path
	
	def txt(self, msg, path):
		f = open(path)
		doc2 = docx.Document()

		for line in f.readlines():
			doc2para = doc2.add_paragraph()
			for symbol in line:
				self.add_text(doc2para, symbol)

		path = os.path.abspath(msg.document.file_name.replace('.txt', '.docx'))
		self.save_file(doc2, path) 
		return path
	
	def text(self, msg):
		doc2 = docx.Document()

		doc2para = doc2.add_paragraph()
		for symbol in msg.text:
			self.add_text(doc2para, symbol)
		
		path = os.path.abspath(f'{msg.chat.id}_{msg.message_id}.docx')
		self.save_file(doc2, path)
		return path
	
	def add_text(self, para, text):
		a 		= random.randint(1,4)
		run 	= para.add_run(text)
		f 		= run.font	
		f.name	= 'Ебать'

		if   a == 1: f.bold, f.italic = False, False
		elif a == 2: f.bold, f.italic =  True, False
		elif a == 3: f.bold, f.italic = False, True
		elif a == 4: f.bold, f.italic =  True, True
	
	def save_file(self, doc, path):
		Mm = docx.shared.Mm
		section 								= doc.sections[0]
		section.page_height 		= Mm(205)
		section.page_width 			= Mm(165)
		section.left_margin 		= Mm(25)
		section.right_margin 		= Mm(10)
		section.top_margin 			= Mm(10)
		section.bottom_margin 	= Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)
		doc.save(path) 